from docx import Document
import re

BLACK = "████████"

def redact_word(input_path, output_path):
    doc = Document(input_path)

    for para in doc.paragraphs:
        text = para.text
        original = text

        # 🔹 Email addresses
        text = re.sub(
            r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+",
            BLACK,
            text
        )

        # 🔹 Phone numbers (10 digits)
        text = re.sub(
            r"\b\d{10}\b",
            BLACK,
            text
        )

        # 🔹 ID numbers (long numbers)
        text = re.sub(
            r"\b\d{4,}\b",
            BLACK,
            text
        )

        # 🔹 Names (capitalized words, 1–3 words)
        text = re.sub(
            r"\b([A-Z][a-z]+(?:\s[A-Z][a-z]+){0,2})\b",
            BLACK,
            text
        )

        # 🔹 Sensitive keywords
        text = re.sub(
            r"\b(Aadhaar|PAN|Account|ATM|Password|Address|IFSC)\b",
            BLACK,
            text,
            flags=re.IGNORECASE
        )

        # Apply changes only if needed
        if text != original:
            para.clear()
            para.add_run(text)

    doc.save(output_path)
